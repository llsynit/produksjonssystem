#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import shutil
import subprocess
import sys
import tempfile
import traceback
import zipfile
from pathlib import Path
import datetime


import re
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from lxml import etree as ElementTree


from core.pipeline import Pipeline
from core.utils.epub import Epub
from core.utils.xslt import Xslt
from core.utils.filesystem import Filesystem

from core.api_queue_worker import add_task
from core.api_worker import ApiWorker



if sys.version_info[0] != 3 or sys.version_info[1] < 5:
    print("# This script requires Python version 3.5+")
    sys.exit(1)

class NLBpubToDocx(Pipeline):
    uid = "nlbpub-to-docx"
    title = "NLBPUB til DOCX"
    labels = ["e-bok", "Statped"]
    publication_format = "XHTML"
    expected_processing_time = 370

    attributes = {
            "edition": "docx",
            "stage": "utgaveutdocx",
        }
    message =""

    xslt_dir = os.path.realpath(os.path.join(os.path.dirname(os.path.realpath(__file__)), "..", "xslt"))

    def on_book_deleted(self):
        self.utils.report.info("Slettet bok i mappa: " + self.book['name'])
        self.utils.report.title = self.title + " EPUB master slettet: " + self.book['name']
        return True

    def on_book_modified(self):
        self.utils.report.info("Endret bok i mappa: " + self.book['name'])
        return self.on_book()

    def on_book_created(self):
        self.utils.report.info("Ny bok i mappa: " + self.book['name'])
        return self.on_book()

    def on_book(self):
        self.utils.report.attachment(None, self.book["source"], "DEBUG")
        epub = Epub(self.utils.report, self.book["source"])

        epubTitle = ""
        try:
            epubTitle = " (" + epub.meta("dc:title") + ") "
        except Exception:
            pass

        # sjekk at dette er en EPUB
        if not epub.isepub():
            self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎"
            message = self.title + ": " + self.book["name"] + " feilet 😭👎"
            ApiWorker.notify(epub.identifier(), "error", message)
            return False

        if not epub.identifier():
            self.utils.report.error(self.book["name"] + ": Klarte ikke å bestemme boknummer basert på dc:identifier.")
            self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎"
            message = self.title + ": " + self.book["name"] + " feilet 😭👎"
            ApiWorker.notify(epub.identifier(), "error", message)
            return False

        # language must be exctracted from epub or else docx default language (nb) wil be used in the converted file
        #language = ""
        dummy = ""
        try:
           #language = " (" + epub.meta("dc:language") + ") "
           #language = epub.meta("dc:language")
            dummy = epub.meta("dc:language")

        except Exception:
            pass

        # ---------- lag en kopi av EPUBen ----------

        temp_epubdir_obj = tempfile.TemporaryDirectory()
        temp_epubdir = temp_epubdir_obj.name
        Filesystem.copy(self.utils.report, self.book["source"], temp_epubdir)
        temp_epub = Epub(self, temp_epubdir)

        opf_path = temp_epub.opf_path()
        if not opf_path:
            self.utils.report.error(self.book["name"] + ": Klarte ikke å finne OPF-fila i EPUBen.")
            self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            message = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            ApiWorker.notify(epub.identifier(), "error", message)
            return False
        opf_path = os.path.join(temp_epubdir, opf_path)
        opf_xml = ElementTree.parse(opf_path).getroot()

        html_file = opf_xml.xpath("/*/*[local-name()='manifest']/*[@id = /*/*[local-name()='spine']/*[1]/@idref]/@href")
        html_file = html_file[0] if html_file else None
        if not html_file:
            self.utils.report.error(self.book["name"] + ": Klarte ikke å finne HTML-fila i OPFen.")
            self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            message = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            ApiWorker.notify(epub.identifier(), "error", message)
            return False
        html_file = os.path.join(os.path.dirname(opf_path), html_file)
        if not os.path.isfile(html_file):
            self.utils.report.error(self.book["name"] + ": Klarte ikke å finne HTML-fila.")
            self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            message = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            ApiWorker.notify(epub.identifier(), "error", message)
            return False

        temp_xml_file_obj = tempfile.NamedTemporaryFile()
        temp_xml_file = temp_xml_file_obj.name

        self.utils.report.info("Konverterer fra ASCIIMath til norsk punktnotasjon…")
        xslt = Xslt(self,
                    stylesheet=os.path.join(Xslt.xslt_dir, NLBpubToDocx.uid, "nordic-asciimath-epub.xsl"),
                    source=html_file,
                    target=temp_xml_file)
        if not xslt.success:
            return False
        shutil.copy(temp_xml_file, html_file)


                # ---------- konverter HTML-fila til DOCX ----------

        temp_docxdir_obj = tempfile.TemporaryDirectory()
        temp_docxdir = temp_docxdir_obj.name

        try:
            self.utils.report.info("Konverterer fra XHTML til DOCX...")
            process = self.utils.filesystem.run([
                "/usr/bin/ebook-convert",
                html_file,
                os.path.join(temp_docxdir, epub.identifier() + "_calibre.docx"),
                "--chapter=/",
                "--chapter-mark=none",
                "--page-breaks-before=/",
                "--no-chapters-in-toc",
                "--toc-threshold=0",
                "--docx-page-size=a4",
                # "--linearize-tables",
                "--extra-css=" + os.path.join(Xslt.xslt_dir, self.uid, 'extra.css'),

                # NOTE: microsoft fonts must be installed:
                # sudo apt-get install ttf-mscorefonts-installer
                "--embed-font-family=Verdana",

                "--docx-page-margin-top=42",
                "--docx-page-margin-bottom=42",
                "--docx-page-margin-left=70",
                "--docx-page-margin-right=56",
                #"--language="+epub.meta('dc:language'),
                #("--language=" + language) if language else "", test av language
                "--base-font-size=13",
                #"--remove-paragraph-spacing",
                #"--remove-paragraph-spacing-indent-size=-1",
                "--font-size-mapping=13,13,13,13,13,13,13,13"
            ])

            if process.returncode == 0:
                self.utils.report.info("Boken ble konvertert.")

# -------------  script from kvile ---------------
                document = Document(os.path.join(temp_docxdir, epub.identifier() + "_calibre.docx"))
                emptyParagraph = False
                normalParagraph = "Normal"
                normalParagraphNoIndent = "NormalNoIndent"
                headingIndent = Cm(1.25)
                fontSize = Pt(13)
                # ny kode 2021-01-20
                #folder = os.path.join(temp_docxdir)

                folder = Path(temp_docxdir)
                # slutt ny kode

                #self.utils.report.info("Folder: "+folder)

                def zipdir(src, dst, zip_name):
                    os.chdir(dst)
                    ziph = zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED)
                    for root, dirs, files in os.walk(src):
                        for file in files:
                            ziph.write(os.path.join(root, file), arcname=os.path.join(root.replace(src, ""), file))
                    ziph.close()

                def writeFile(txt, dst):
                    tempFile = open(folder / dst,"w+")
                    tempFile.write(txt)
                    tempFile.close()

                def delete_paragraph(paragraph):
                   # self.utils.report.info("Delete paragraph: ")
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None

                def delete_element(element):
                    element.getparent().remove(element)
                    element._element = None

                def isPageNr(element):
                    return re.match(r"^--- \d+( til \d+)?$", element.text)

                def isEmpty(element):
                    return element.text.strip() == "" and not(element._element.xpath(".//w:hyperlink"))

                indent = Cm(0.44)
                hangingIndentList = Cm(0.63)
                document.styles[normalParagraph].font.size = fontSize
                document.styles[normalParagraph].paragraph_format.first_line_indent = indent
                styleNoIndent = document.styles.add_style('NormalNoIndent', WD_STYLE_TYPE.PARAGRAPH)
                styleNoIndent.base_style = document.styles[normalParagraph]
                document.styles[normalParagraphNoIndent].paragraph_format.first_line_indent = Cm(0)

                # set style to normal for regular paragraphs, set keep_with_next to false, remove multiple empty paragraphs, and remove empty p after page nr or heading
                for paragraph in document.paragraphs:
                    # deleting empty text-elements
                    emptyTextElementList = document.element.xpath("//w:t[. = '']")
                    for emptyTextElement in emptyTextElementList:
                        delete_element(emptyTextElement)
                    paragraph.paragraph_format.keep_with_next = None
                    if re.match("Para 0[1-9]|[0-9] Block|Para [0-9]", paragraph.style.name): # and paragraph.style.font.underline != True: #TODO 08.12.23 this line and line 285 commented out -- probable effect will be that link will not look like a link.
                        paragraph.style = normalParagraph
                    if isEmpty(paragraph) or isPageNr(paragraph) or paragraph.style.name[0:7] == "Heading": # if empty p or page nr or heading
                    # if len(paragraph.text) <= 1 or isPageNr(paragraph) or paragraph.style.name[0:7] == "Heading": # if empty p or page nr or heading
                        paragraph.text = paragraph.text.strip() # added 26.09.23. docx 012
                        #paragraph.text = re.sub(r"^\s(.*)", r"\1", paragraph.text)  #remove space at beginning av p <p><a>statped.no</a></p> <p><a></a>
                       # self.utils.report.info("Paragraph.text <= 1 ")
                        # edited 2023-09-20
                        if isEmpty(paragraph) and emptyParagraph: #if last p also was empty or page nr
                        # if len(paragraph.text) == 0 and emptyParagraph: #if last p also was empty or page nr
                    #        self.utils.report.info("Paragraph.text == 0 ")
                            delete_paragraph(paragraph)
                        emptyParagraph = True
                    else:
                        emptyParagraph = False
                        if re.match(r"^\s*STATPED_DUMMYTEXT_LI_OL\s*$", paragraph.text):
                            paragraph.text = ""
                # no indent after Heading, page-nr, or paragraphs starting with "Bilde: ", paragraphs in only bold (text=^_[^_]*_$) and the paragraph after p in only bold, or on empty p.
                removeIndent = False
                for paragraph in document.paragraphs:
                    #remove space at beginning of line after <br/>
                    spaceAfterBreakList = paragraph._element.xpath(r'w:r/w:br[@w:clear="none"]/following::w:t[@xml:space="preserve"][1]')
                    if len(spaceAfterBreakList) > 0:
                        for spaceAfterBreakElement in spaceAfterBreakList:
                            if re.match('^ ', spaceAfterBreakElement.text) and not(spaceAfterBreakElement.xpath(r'preceding-sibling::*[1][self::w:t]')):
                                spaceAfterBreakElement.text = re.sub(r"^ ", r"", spaceAfterBreakElement.text)
                    #remove break before paragraph end
                        breakBeforeParagraphEndList = paragraph._element.xpath(r'w:r[last()]/w:br[@w:clear="none" and not(following-sibling::*)]')
                        if len(breakBeforeParagraphEndList) > 0:
                            delete_element(breakBeforeParagraphEndList[0])

                    t = paragraph.text.strip()
                    if re.match(r"^Bilde: |^Forklaring: |^--- \d+ til |^_[^_]*_$|^STATPED_DUMMYTEXT_LIST_UNSTYLED|^STATPED_DUMMYTEXT_P_BEFORE_DL", t) or ((removeIndent or isEmpty(paragraph)) and paragraph.style.name == "Normal"):
                        paragraph.style = normalParagraphNoIndent
                    # Remove dummy-text and set hengemarg
                    if re.match(r"^(STATPED_DUMMYTEXT_LIST_UNSTYLED|STATPED_DUMMYTEXT_DL)", paragraph.text):
                        paragraph.paragraph_format.left_indent = hangingIndentList #Pt(0)
                        paragraph.paragraph_format.first_line_indent = -hangingIndentList #Pt(-20)
                    if re.match(r"^STATPED_DUMMYTEXT", paragraph.text):
                        paragraph.text = re.sub(r"^(STATPED_DUMMYTEXT_LIST_UNSTYLED|STATPED_DUMMYTEXT_DL|STATPED_DUMMYTEXT_P_BEFORE_DL)", "", paragraph.text)
                    if len(t) == 0 or paragraph.style.name[0:7] == "Heading" or re.match(r"^--- \d+ til |^_[^_]*_$", t):
                        removeIndent = True
                    else:
                        removeIndent = False

                # remove bold from Headings.
                paraStylesWithoutBoldOrUnderline = [] #list of all para-styles without underline or bold
                paraStylesWithoutUnderline = [] #list of all para-styles without underline
                for style in document.styles:
                    if style.name[0:7] == "Heading":
                        style.font.bold = None
                        style.paragraph_format.left_indent = headingIndent #Pt(0)
                        style.paragraph_format.first_line_indent = -headingIndent #Pt(-20)
                        style.paragraph_format.space_before = Pt(0)
                        style.paragraph_format.space_after = Pt(0)
                        style_element = style._element
                        spacing = style_element.xpath(r'w:pPr/w:spacing')[0]
                        spacing.set(qn('w:beforeLines'), "0")
                        spacing.set(qn('w:afterLines'), "0")
                    if style.name[0:5] == "Para ":
                        #if style.font.underline != True: commented out 08.12.23
                            paraStylesWithoutUnderline.append(style.name)
                            if style.font.bold != True:
                                paraStylesWithoutBoldOrUnderline.append(style.name)

                # find all para-styles with wanted properties in tables and change style
                paraStylesInTables = []
                #for paraStyleWithoutBoldOrUnderline in paraStylesWithoutBoldOrUnderline:
                for paraStyleWithoutUnderline in paraStylesWithoutUnderline:
                    for element in document.element.xpath("//w:tbl//w:p//w:pStyle[@w:val = '" + paraStyleWithoutUnderline + "']"):
                        paraStylesInTables.append(element)
                for paraStyleInTables in paraStylesInTables:
                    paraStyleInTables.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] =  normalParagraphNoIndent # or normalParagraph

                # uncomment if you want to modify first p in a cell
                # firstParaStylesInTables = []
                # for paraStyleWithoutBoldOrUnderline in paraStylesWithoutBoldOrUnderline:
                #     for element in document.element.xpath("//w:tc//w:p[position()=1]//w:pStyle[@w:val = '" + normalParagraph + "']"):
                #         firstParaStylesInTables.append(element)
                # for paraStyleInTables in firstParaStylesInTables:
                #     paraStyleInTables.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] = normalParagraphNoIndent


                # tables missing required <w:tblGrid>, so throws: docx.oxml.exceptions.InvalidXmlError: required ``<w:tblGrid>`` child element not present
                #from docx.table import _Cell, Table
                #from docx.oxml.text.paragraph import CT_P
                # for row in table.columns:
                #     try:
                #         for cell in row.cells:
                #             firstP = True
                #             for p in cell.paragraphs:
                #                 if p.style.font.underline != True and re.match(r"^Para | Block",p.style.name):
                #                     if firstP:
                #                         p.style = "NormalNoIndent"
                #                         firstP = False
                #                     else:
                #                         p.style = "Normal"
                #     except Exception as e:
                #         pass


                document.save(os.path.join(temp_docxdir, epub.identifier() + ".docx"))
                self.utils.report.info("Temp-fil ble lagret: "+os.path.join(temp_docxdir, epub.identifier() + ".docx"))

                wordFile = os.path.join(temp_docxdir, epub.identifier() + ".docx")

                zipDocument = zipfile.ZipFile((folder / wordFile))
                tempFolder = "temp"
                zipDocument.extractall(folder/ tempFolder)
                zipDocument.close()
                zippedFile = tempFolder + "/word/numbering.xml"
                xmlFile = open((folder / zippedFile), 'r+')
                xmlText = xmlFile.read()
                xmlText = re.sub(r'w:left="1152"', r'w:left="360"', xmlText)
                xmlText = re.sub(r'w:left="1512"', r'w:left="720"', xmlText)
                xmlText = re.sub(r'w:left="1872"', r'w:left="1080"', xmlText)
                xmlText = re.sub(r'<w:numFmt w:val="lowerLetter"/><w:lvlText w:val="%([1-9])\."/>', r'<w:numFmt w:val="lowerLetter"/><w:lvlText w:val="%\1)"/>', xmlText) # a. as a) in lists
                #xmlText = re.sub(r'<w:lvlText w:val="%(1|2)\."/>', r'<w:lvlText w:val="%\1)"/>', xmlText) # a. as a), and 1. as 1) in lists

                writeFile(xmlText, zippedFile)
                zipdir(str(folder / tempFolder), str(folder), os.path.join(temp_docxdir, epub.identifier() + ".docx"))

# ---------- end script from kvile -------

            else:
                self.utils.report.error("En feil oppstod ved konvertering til DOCX for " + epub.identifier())
                self.utils.report.debug(traceback.format_stack())
                self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
                message = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
                ApiWorker.notify(epub.identifier(), "error", message)
                return False

        except subprocess.TimeoutExpired:
            self.utils.report.error("Det tok for lang tid å konvertere " + epub.identifier() + " til DOCX, og Calibre-prosessen ble derfor stoppet.")
            self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            message = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            ApiWorker.notify(epub.identifier(), "error", message)
            return False

        except Exception:
            self.utils.report.error("En feil oppstod ved konvertering til DOCX for " + epub.identifier())
            self.utils.report.info(traceback.format_exc(), preformatted=True)
            self.utils.report.title = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            message = self.title + ": " + self.book["name"] + " feilet 😭👎" + epubTitle
            ApiWorker.notify(epub.identifier(), "error", message)
            return False

        archived_path, stored = self.utils.filesystem.storeBook(temp_docxdir, epub.identifier())
        date_modified = datetime.datetime.now().timestamp()
        self.utils.report.attachment(None, archived_path, "DEBUG")
        self.utils.report.title = self.title + ": " + epub.identifier() + " ble konvertert 👍😄" + epubTitle
        dt_m = datetime.datetime.fromtimestamp(date_modified)
        archived_time = dt_m.isoformat()
        self.attributes["date_modified"] = archived_time
        add_task(self.uid,epub.identifier(), self.attributes)
        message = self.title + ": " + self.book["name"] + " ble konvertert 👍😄" + epubTitle
        ApiWorker.notify(epub.identifier(), "success", message)
        return True

if __name__ == "__main__":
    NLBpubToDocx().run()
